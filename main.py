import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import mysqlgen
import os.path

my_path = os.path.abspath(os.path.dirname(__file__))

title = 'ASSIGNMENT'
document_name = title + '.docx'
document_path = os.path.join(my_path, document_name)
document = docx.Document()

titleKey = 'TitleStyle'
headingKey = 'HeadingStyle'
textKey = 'DefaultTextStyle'
codeKey = 'CodeStyle'

def create_styles():
    styles = document.styles

    
    titleStyle = styles.add_style(titleKey, WD_STYLE_TYPE.PARAGRAPH)
    font = titleStyle.font
    font.size = Pt(16)
    font.bold = True
    font.underline = True
    font.name = 'Arial'
    titleStyle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


    headingStyle = styles.add_style(headingKey, WD_STYLE_TYPE.PARAGRAPH)
    font = headingStyle.font
    font.size = Pt(12)
    font.bold = True
    font.underline = True
    font.name = 'Arial'


    codeStyle = styles.add_style(codeKey, WD_STYLE_TYPE.CHARACTER)
    font = codeStyle.font
    font.size = Pt(8)
    font.name = 'Consolas'


    textStyle = styles.add_style(textKey, WD_STYLE_TYPE.CHARACTER)
    font = textStyle.font
    font.size = Pt(11)
    font.name = 'Arial'

    document.save(document_path)


def create_document(ass_num, data):
    document.add_paragraph(title + ' #' + ass_num + '\n', style = titleKey)

    for heading in data:
        document.add_paragraph(heading, style = headingKey)

        style = codeKey if heading == 'QUERY' or heading == 'OUTPUT' else textKey
        r = document.add_paragraph().add_run(data[heading]  + '\n', style = style)
        if heading == 'QUERY': r.font.size = Pt(11)

    document.add_page_break()
    document.save(document_name)


def read_file():
    with open('queries.txt') as f:
        line = f.readline().strip()
        while line != 'end':
            aim = f.readline().strip()
            query = f.readline().strip()

            undo = f.readline().strip()
            if undo[0] == '#':
                table = mysqlgen.execute_query(query, undo[1:])
                undo = f.readline()
            else:
                table = mysqlgen.execute_query(query)

            data_dict = {
                'AIM' : aim,
                'QUERY' : query,
                'OUTPUT' : table,
                'RESULT' : 'Query executed successfully.'
            }

            create_document(line[:-1], data_dict)
            print('Done', line)
            line = undo 
    mysqlgen.closeDb()   

create_styles()
read_file()


